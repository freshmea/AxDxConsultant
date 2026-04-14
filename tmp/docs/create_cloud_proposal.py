from pathlib import Path

import win32com.client  # type: ignore
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
import pypdfium2 as pdfium


BASE = Path(r"C:\Users\Administrator\dxAx")
DOCX_PATH = BASE / "output" / "doc" / "sample.docx"
PDF_PATH = BASE / "output" / "pdf" / "sample.pdf"
PNG_PREFIX = BASE / "tmp" / "pdfs" / "sample_page"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def shade(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, size=10.5, bold=False, color="222222"):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)

# Header block
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
header_table.columns[0].width = Cm(12.9)
header_table.columns[1].width = Cm(5.5)

left = header_table.cell(0, 0)
p = left.paragraphs[0]
r = p.add_run("내년도 클라우드 도입 추진 제안서")
set_font(r, size=14.5, bold=True, color="3A3A3A")
p.space_after = Pt(2)

p2 = left.add_paragraph("company name")
r2 = p2.runs[0]
set_font(r2, size=12, bold=True, color="4B4B4B")
p2.space_after = Pt(4)

p3 = left.add_paragraph("전략기획실 AX추진TF")
set_font(p3.runs[0], size=10, color="666666")

right = header_table.cell(0, 1)
approval = right.add_table(rows=2, cols=3)
approval.autofit = False
for i in range(3):
    approval.columns[i].width = Cm(1.78)
approval.rows[0].height = Cm(0.55)
approval.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
approval.rows[1].height = Cm(1.0)
approval.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
labels = ["담당", "팀장", "사장"]
for i, label in enumerate(labels):
    cp = approval.cell(0, i).paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cp.add_run(label)
    set_font(rr, size=9.5, bold=True, color="4A4A4A")
    shade(approval.cell(0, i), "E9E2D6")
    approval.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
for i in range(3):
    cp = approval.cell(1, i).paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run("")
    approval.cell(1, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for row in approval.rows:
    for cell in row.cells:
        for edge in ("left", "top", "right", "bottom"):
            set_cell_border(cell, **{edge: {"val": "single", "sz": 4, "color": "B7A58E"}})

doc.add_paragraph("")

# Main form table
t = doc.add_table(rows=8, cols=6)
t.autofit = False
col_widths = [Cm(2.2), Cm(4.0), Cm(2.1), Cm(3.0), Cm(2.2), Cm(3.1)]
for i, w in enumerate(col_widths):
    t.columns[i].width = w

def fill_row_label(row, idx, text):
    c = t.cell(row, idx)
    p = c.paragraphs[0]
    rr = p.add_run(text)
    set_font(rr, size=10, bold=True, color="404040")
    shade(c, "F1EFEC")
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def fill_row_text(row, idx, text, bold=False):
    c = t.cell(row, idx)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = p.add_run(text)
    set_font(rr, size=10, bold=bold, color="2E2E2E")
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

fill_row_label(0, 0, "일시")
fill_row_text(0, 1, "2026. 04. 14")
fill_row_label(0, 2, "작성자")
fill_row_text(0, 3, "AX 추진TF")
fill_row_label(0, 4, "참여인원")
fill_row_text(0, 5, "6명")

fill_row_label(1, 0, "제안명")
t.cell(1, 1).merge(t.cell(1, 5))
fill_row_text(1, 1, "클라우드 전환 기반 운영효율화 및 비용구조 개선안", bold=True)

fill_row_label(2, 0, "현행 한계")
t.cell(2, 1).merge(t.cell(2, 5))
fill_row_text(
    2,
    1,
    "• 현장 중심 계측사업 특성상 장비/서버 이중 운영으로 장애 대응 리드타임이 길고, "
    "• 온프레미스 증설 시 선투자(CAPEX) 부담이 커 수요 변동 대응이 지연되며, "
    "• 유지보수 인력의 반복 점검 업무 비중이 높아 운영 효율이 저하됨.",
)

fill_row_label(3, 0, "도입 필요성")
t.cell(3, 1).merge(t.cell(3, 5))
fill_row_text(
    3,
    1,
    "클라우드 기반으로 데이터 수집/저장/분석/백업을 통합하면, 현장 인프라 의존도를 낮추고 "
    "수요 기반 탄력 운영이 가능해져 장애복구, 확장, 보안통제의 표준화가 가능함.",
)

fill_row_label(4, 0, "기대효과")
t.cell(4, 1).merge(t.cell(4, 5))
fill_row_text(
    4,
    1,
    "비용: 서버 교체·유지비/전력/백업매체 절감으로 연간 약 1,000만 원 절감 가능\n"
    "운영: 무중단 배포·원격 모니터링 자동화로 장애 대응시간 30% 단축\n"
    "확장성: 신규 현장 추가 시 인프라 리드타임 4주 → 3일 수준 단축\n"
    "보안: 중앙 통합계정·접근통제·로그감사 체계 고도화\n"
    "유지관리: 패치/백업 정책 자동화로 운영 인력의 고부가 업무 전환",
)

fill_row_label(5, 0, "적용 방향")
t.cell(5, 1).merge(t.cell(5, 5))
fill_row_text(
    5,
    1,
    "1단계(분기1): 수집 데이터 백업·아카이빙 및 개발/테스트 환경 전환\n"
    "2단계(분기2): 관제 대시보드·분석 워크로드 이전, 현장 게이트웨이 표준화\n"
    "3단계(분기3~4): 핵심 운영계 점진 전환 및 FinOps 체계 정착",
)

fill_row_label(6, 0, "결론/제안")
t.cell(6, 1).merge(t.cell(6, 5))
fill_row_text(
    6,
    1,
    "내년도 예산/조직계획에 클라우드 전환 과제를 반영해 선제 추진 필요. "
    "초기 6개월은 하이브리드로 리스크를 통제하고, KPI(비용절감, MTTR, 신규현장 리드타임)로 성과를 관리할 것을 제안함.",
    bold=True,
)

fill_row_label(7, 0, "예산 영향")
t.cell(7, 1).merge(t.cell(7, 5))
fill_row_text(7, 1, "예시) 기존 연 4,000만 원 운영비 → 전환 후 연 3,000만 원 수준(약 25% 절감, 연 1,000만 원)", bold=True)

for row in t.rows:
    for cell in row.cells:
        for edge in ("left", "top", "right", "bottom"):
            set_cell_border(cell, **{edge: {"val": "single", "sz": 4, "color": "B7A58E"}})

for p in doc.paragraphs:
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25

DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(DOCX_PATH))

# Convert DOCX to PDF using Word COM
word = win32com.client.DispatchEx("Word.Application")
word.Visible = False
word.DisplayAlerts = 0
try:
    doc_obj = word.Documents.Open(str(DOCX_PATH), ReadOnly=True)
    doc_obj.ExportAsFixedFormat(OutputFileName=str(PDF_PATH), ExportFormat=17)
    doc_obj.Close(False)
finally:
    word.Quit()

# Render first page preview PNG for visual QA
pdf = pdfium.PdfDocument(str(PDF_PATH))
page = pdf[0]
bitmap = page.render(scale=2.0)
pil_img = bitmap.to_pil()
pil_img.save(str(PNG_PREFIX) + "_1.png")
pdf.close()

print(f"DOCX: {DOCX_PATH}")
print(f"PDF: {PDF_PATH}")
print(f"PNG: {PNG_PREFIX}_1.png")
